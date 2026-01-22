"""Module for reading content from various file types."""

from pathlib import Path
from typing import Dict, Any

from config import SUPPORTED_EXTENSIONS, MAX_FILE_SIZE_MB


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get information about a file.
    
    Args:
        file_path: Path to the file.
        
    Returns:
        Dictionary with file information.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    stat = path.stat()
    size_mb = stat.st_size / (1024 * 1024)
    extension = path.suffix.lower()
    
    return {
        'name': path.name,
        'extension': extension,
        'type': SUPPORTED_EXTENSIONS.get(extension, 'unknown'),
        'size_bytes': stat.st_size,
        'size_mb': round(size_mb, 2),
        'path': str(path.absolute()),
    }


def validate_file(file_path: str) -> None:
    """
    Validate a file before reading.
    
    Args:
        file_path: Path to the file to validate.
        
    Raises:
        FileNotFoundError: If file doesn't exist.
        ValueError: If file type not supported or file too large.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ', '.join(SUPPORTED_EXTENSIONS.keys())
        raise ValueError(f"Unsupported file type: {extension}\nSupported types: {supported}")
    
    # Check file size
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(f"File too large: {size_mb:.1f}MB (max: {MAX_FILE_SIZE_MB}MB)")


def read_text_file(file_path: str) -> str:
    """Read content from a plain text file."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    # Fallback: read with errors ignored
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def read_pdf_file(file_path: str) -> str:
    """Read content from a PDF file."""
    try:
        import pypdf
    except ImportError:
        raise ImportError(
            "pypdf is required to read PDF files.\n"
            "Install with: pip install pypdf"
        )
    
    reader = pypdf.PdfReader(file_path)
    text_content = []
    
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_content.append(f"--- Page {i + 1} ---\n{page_text}")
    
    return '\n\n'.join(text_content)


def read_docx_file(file_path: str) -> str:
    """Read content from a Word document."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to read DOCX files.\n"
            "Install with: pip install python-docx"
        )
    
    doc = Document(file_path)
    content_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading
            if para.style and para.style.name.startswith('Heading'):
                content_parts.append(f"\n## {text}\n")
            else:
                content_parts.append(text)
    
    # Also extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(' | '.join(cells))
        if table_rows:
            content_parts.append('\n' + '\n'.join(table_rows) + '\n')
    
    return '\n'.join(content_parts)


def read_csv_file(file_path: str) -> str:
    """Read content from a CSV file with better formatting."""
    import csv
    
    with open(file_path, 'r', encoding='utf-8', newline='') as f:
        reader = csv.reader(f)
        rows = list(reader)
    
    if not rows:
        return ""
    
    # Format as a readable table
    result_lines = []
    
    # Header
    if rows:
        header = rows[0]
        result_lines.append("Columns: " + ', '.join(header))
        result_lines.append("-" * 50)
    
    # Data rows (limit for very large files)
    max_rows = 1000
    for i, row in enumerate(rows[1:max_rows + 1], 1):
        if header:
            row_data = [f"{h}: {v}" for h, v in zip(header, row)]
            result_lines.append(f"Row {i}: " + ' | '.join(row_data))
        else:
            result_lines.append(f"Row {i}: " + ', '.join(row))
    
    if len(rows) > max_rows + 1:
        result_lines.append(f"\n... and {len(rows) - max_rows - 1} more rows")
    
    return '\n'.join(result_lines)


def read_json_file(file_path: str) -> str:
    """Read content from a JSON file with pretty formatting."""
    import json
    
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Format with indentation for readability
    return json.dumps(data, indent=2, ensure_ascii=False)


def read_xml_file(file_path: str) -> str:
    """Read content from an XML file."""
    import xml.etree.ElementTree as ET
    
    tree = ET.parse(file_path)
    root = tree.getroot()
    
    def extract_text(element, depth=0):
        """Recursively extract text from XML elements."""
        lines = []
        indent = "  " * depth
        
        # Element tag and attributes
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        attrs = ' '.join(f'{k}="{v}"' for k, v in element.attrib.items())
        
        # Element text
        text = element.text.strip() if element.text else ""
        
        if text:
            lines.append(f"{indent}{tag}: {text}")
        elif attrs:
            lines.append(f"{indent}{tag} ({attrs})")
        
        # Children
        for child in element:
            lines.extend(extract_text(child, depth + 1))
        
        return lines
    
    return '\n'.join(extract_text(root))


def read_html_file(file_path: str) -> str:
    """Read and extract text content from an HTML file."""
    import re
    
    content = read_text_file(file_path)
    
    # Remove script and style elements
    content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert some HTML to readable format
    content = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n## \1\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
    
    # Remove remaining HTML tags
    content = re.sub(r'<[^>]+>', '', content)
    
    # Clean up whitespace
    content = re.sub(r'\n\s*\n', '\n\n', content)
    content = re.sub(r' +', ' ', content)
    
    return content.strip()


def read_file(file_path: str) -> str:
    """
    Read content from a file based on its extension.
    
    Args:
        file_path: Path to the file to read.
        
    Returns:
        The content of the file as a string.
        
    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the file does not exist.
    """
    # Validate file first
    validate_file(file_path)
    
    path = Path(file_path)
    extension = path.suffix.lower()
    file_type = SUPPORTED_EXTENSIONS[extension]
    
    # Route to appropriate reader
    readers = {
        'pdf': read_pdf_file,
        'docx': read_docx_file,
        'csv': read_csv_file,
        'json': read_json_file,
        'xml': read_xml_file,
        'html': read_html_file,
    }
    
    reader = readers.get(file_type, read_text_file)
    return reader(file_path)
